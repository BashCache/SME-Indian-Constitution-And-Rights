"""
DOCX Generator: Generates Microsoft Word documents from text content.
"""

from typing import Optional
from docx import Document
from docx.shared import Inches as DocxInches, Pt as DocxPt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .document_generator import DocumentGenerator

class DocxGenerator(DocumentGenerator):
    """Generator for DOCX (Microsoft Word) documents"""
    
    def generate(self, content: str, output_path: str, title: Optional[str] = None,
                author: Optional[str] = None, subject: Optional[str] = None) -> str:
        """Generate a DOCX document from text content"""
        doc = Document()
        
        # Set document properties
        core_props = doc.core_properties
        if title:
            core_props.title = title
        if author:
            core_props.author = author
        if subject:
            core_props.subject = subject
        
        # Add title if provided
        if title:
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Add spacing
        
        # Split content into paragraphs
        paragraphs = self.split_into_paragraphs(content)
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            # Check if it's a heading
            if self.is_heading(para_text):
                level, heading_text = self.extract_heading_level(para_text)
                doc.add_heading(heading_text, level=level)
            else:
                # Regular paragraph
                para = doc.add_paragraph(para_text)
                para_format = para.paragraph_format
                para_format.space_after = DocxPt(12)
                para_format.first_line_indent = DocxInches(0.5)
        
        # Save document
        doc.save(output_path)
        return output_path

